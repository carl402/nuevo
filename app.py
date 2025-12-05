import os
import io
from datetime import datetime, timedelta
from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from models import db, User, Project, Variable, Report, SimulationResult
from dotenv import load_dotenv
from utils import simulate_monte_carlo
from weasyprint import HTML, CSS
from docx import Document
from apscheduler.schedulers.background import BackgroundScheduler

load_dotenv()

APP_DIR = os.path.dirname(os.path.abspath(__file__))
DATABASE_URL = os.environ.get("DATABASE_URL") or "postgresql://postgres:montecarlo@db.cowlhqyvwlreakopucki.supabase.co:5432/postgres"
SECRET_KEY = os.environ.get("SECRET_KEY") or "change-me"

def create_app():
    app = Flask(__name__)
    app.config["SQLALCHEMY_DATABASE_URI"] = DATABASE_URL
    app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
    app.config["SECRET_KEY"] = SECRET_KEY

    db.init_app(app)

    login_manager = LoginManager()
    login_manager.login_view = "login"
    login_manager.init_app(app)

    @login_manager.user_loader
    def load_user(user_id):
        return User.query.get(int(user_id))

    # --- Routes ---
    @app.route("/")
    def index():
        if current_user.is_authenticated:
            return redirect(url_for("dashboard"))
        return redirect(url_for("login"))

    @app.route("/login", methods=["GET", "POST"])
    def login():
        if request.method == "POST":
            email = request.form.get("email")
            password = request.form.get("password")
            user = User.query.filter_by(email=email).first()
            if user and user.check_password(password):
                login_user(user)
                return redirect(url_for("dashboard"))
            flash("Credenciales inválidas", "danger")
        return render_template("login.html")

    @app.route("/logout")
    @login_required
    def logout():
        logout_user()
        return redirect(url_for("login"))

    @app.route("/dashboard")
    @login_required
    def dashboard():
        return redirect(url_for("projects_list"))

    @app.route("/simulaciones")
    @login_required
    def simulaciones():
        recent_reports = Report.query.order_by(Report.created_at.desc()).limit(10).all()
        return render_template("simulaciones.html", reports=recent_reports)

    # Projects
    @app.route("/projects")
    @login_required
    def projects_list():
        q = request.args.get("q")
        if q:
            projects = Project.query.filter(Project.name.ilike(f"%{q}%") ).all()
        else:
            projects = Project.query.all()
        return render_template("projects.html", projects=projects, q=q)

    @app.route("/projects/create", methods=["POST"]) 
    @login_required
    def projects_create():
        name = request.form.get("name")
        description = request.form.get("description")
        status = request.form.get("status") or "Activo"
        if not name:
            flash("El nombre es requerido", "warning")
            return redirect(url_for("projects_list"))
        p = Project(name=name, description=description, user_id=current_user.id)
        db.session.add(p)
        db.session.commit()
        flash("Proyecto creado", "success")
        return redirect(url_for("projects_list") + f"?open_config={p.id}")

    @app.route("/projects/<int:project_id>/delete", methods=["POST"])
    @login_required
    def projects_delete(project_id):
        p = Project.query.get_or_404(project_id)
        db.session.delete(p)
        db.session.commit()
        flash("Proyecto eliminado", "info")
        return redirect(url_for("projects_list"))

    # Variables: add/remove
    @app.route("/projects/<int:project_id>/variables/add", methods=["POST"])
    @login_required
    def variable_add(project_id):
        p = Project.query.get_or_404(project_id)
        name = request.form.get("var_name")
        dist = request.form.get("var_dist").lower()
        
        var = Variable(project=p, name=name, distribution=dist)
        
        if dist == "normal":
            var.mean = request.form.get("mean", 50)
            var.std_dev = request.form.get("std", 15)
        elif dist == "uniform":
            var.min_value = request.form.get("low", 0)
            var.max_value = request.form.get("high", 100)
        elif dist == "triangular":
            var.min_value = request.form.get("left", 10)
            var.mode_value = request.form.get("mode", 50)
            var.max_value = request.form.get("right", 90)

        db.session.add(var)
        db.session.commit()
        flash("Variable añadida", "success")
        return redirect(url_for("projects_list"))

    @app.route("/projects/<int:project_id>/simulate", methods=["POST"])
    @login_required
    def project_simulate(project_id):
        p = Project.query.get_or_404(project_id)
        iterations = int(request.form.get("iterations", 1000))
        seed = request.form.get("seed") or None
        variables = []
        for v in p.variables:
            variables.append({"name": v.name, "distribution": v.distribution, "params": v.params})

        results = simulate_monte_carlo(variables, iterations=iterations, seed=seed)

        r = Report(project=p, name=f"Reporte {p.name} {datetime.utcnow().isoformat()}")
        db.session.add(r)
        db.session.flush()
        
        # Guardar resultados en simulation_results
        summary = results.get("summary", {})
        sim_result = SimulationResult(
            report_id=r.id,
            mean_value=summary.get("mean"),
            std_dev=summary.get("std"),
            variance_value=summary.get("variance"),
            min_value=summary.get("min"),
            max_value=summary.get("max"),
            percentile_5=summary.get("percentile_5"),
            percentile_95=summary.get("percentile_95"),
            median_value=summary.get("median")
        )
        db.session.add(sim_result)
        
        # Guardar gráficos en la sesión para mostrar en el reporte
        from flask import session
        session[f'charts_{r.id}'] = results.get('charts', {})
        
        db.session.commit()
        flash("Simulación ejecutada y reporte guardado", "success")
        return redirect(url_for("report_view", report_id=r.id))

    # Reports
    @app.route("/reports")
    @login_required
    def reports_list():
        q = request.args.get("project")
        date_filter = request.args.get("date")
        query = Report.query
        if q:
            query = query.join(Project).filter(Project.name.ilike(f"%{q}%"))
        if date_filter:
            try:
                date_obj = datetime.fromisoformat(date_filter)
                query = query.filter(db.func.date(Report.created_at) == date_obj.date())
            except ValueError:
                pass
        reports = query.order_by(Report.created_at.desc()).all()
        return render_template("reports.html", reports=reports, deleted=[])

    @app.route("/reports/<int:report_id>")
    @login_required
    def report_view(report_id):
        r = Report.query.get_or_404(report_id)
        return render_template("report_view.html", report=r)

    @app.route("/reports/<int:report_id>/delete", methods=["POST"])
    @login_required
    def report_delete(report_id):
        r = Report.query.get_or_404(report_id)
        db.session.delete(r)
        db.session.commit()
        flash("Reporte eliminado", "info")
        return redirect(url_for("reports_list"))

    @app.route("/reports/<int:report_id>/download/pdf")
    @login_required
    def report_download_pdf(report_id):
        r = Report.query.get_or_404(report_id)
        html = render_template("report_pdf.html", report=r)
        try:
            pdf_file = HTML(string=html).write_pdf()
            return send_file(io.BytesIO(pdf_file), mimetype="application/pdf", as_attachment=True, download_name=f"report_{r.id}.pdf")
        except Exception as e:
            flash(f"Error generando PDF: {e}", "danger")
            return redirect(url_for("report_view", report_id=report_id))

    @app.route("/reports/<int:report_id>/download/docx")
    @login_required
    def report_download_docx(report_id):
        r = Report.query.get_or_404(report_id)
        doc = Document()
        doc.add_heading(r.name, level=1)
        doc.add_paragraph(f"Proyecto: {r.project.name}")
        doc.add_paragraph(f"Creado: {r.created_at}")
        doc.add_heading("Resumen", level=2)
        for k, v in r.results.get("summary", {}).items():
            doc.add_paragraph(f"{k}: {v}")
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return send_file(bio, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", as_attachment=True, download_name=f"report_{r.id}.docx")

    # Users management (minimal)
    @app.route("/users")
    @login_required
    def users_list():
        users = User.query.all()
        return render_template("users.html", users=users)

    @app.route("/users/create", methods=["POST"])
    @login_required
    def users_create():
        first = request.form.get("first_name")
        last = request.form.get("last_name")
        email = request.form.get("email")
        password = request.form.get("password")
        role = request.form.get("role") or "user"
        if not (first and email and password):
            flash("Faltan campos requeridos", "warning")
            return redirect(url_for("users_list"))
        u = User(first_name=first, last_name=last, email=email, role=role)
        u.set_password(password)
        db.session.add(u)
        db.session.commit()
        flash("Usuario creado", "success")
        return redirect(url_for("users_list"))

    @app.route("/users/<int:user_id>/edit", methods=["POST"])
    @login_required
    def users_edit(user_id):
        u = User.query.get_or_404(user_id)
        u.first_name = request.form.get("first_name") or u.first_name
        u.last_name = request.form.get("last_name") or u.last_name
        u.email = request.form.get("email") or u.email
        u.role = request.form.get("role") or u.role
        db.session.commit()
        flash("Usuario actualizado", "success")
        return redirect(url_for("users_list"))

    @app.route("/users/<int:user_id>/password", methods=["POST"])
    @login_required
    def users_change_password(user_id):
        u = User.query.get_or_404(user_id)
        new_password = request.form.get("new_password")
        if not new_password:
            flash("La contraseña es requerida", "warning")
            return redirect(url_for("users_list"))
        u.set_password(new_password)
        db.session.commit()
        flash("Contraseña actualizada", "success")
        return redirect(url_for("users_list"))

    @app.route("/users/<int:user_id>/delete", methods=["POST"])
    @login_required
    def users_delete(user_id):
        u = User.query.get_or_404(user_id)
        if u.id == current_user.id:
            flash("No puedes eliminar tu propio usuario", "danger")
            return redirect(url_for("users_list"))
        db.session.delete(u)
        db.session.commit()
        flash("Usuario eliminado", "info")
        return redirect(url_for("users_list"))

    # simple CLI action to init DB
    @app.cli.command("init-db")
    def init_db():
        with app.app_context():
            db.create_all()
            print("DB initialized")



    return app


app = create_app()

if __name__ == "__main__":
    import sys
    # quick init option
    if "--init-db" in sys.argv:
        with app.app_context():
            db.create_all()
            print("DB initialized")
            sys.exit(0)
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
